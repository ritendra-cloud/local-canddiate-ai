from docx import Document
from app.services.resume_import_service import extract_docx, build_profile
def test_extracts_paragraphs_and_tables(tmp_path):
    file=tmp_path/'resume.docx'; doc=Document(); doc.add_paragraph('Jane Doe'); doc.add_paragraph('Skills'); doc.add_paragraph('Python, SQL'); table=doc.add_table(rows=1, cols=2); table.cell(0,0).text='Role'; table.cell(0,1).text='Engineer'; doc.save(file)
    before=file.read_bytes(); lines=extract_docx(file)
    assert 'Jane Doe' in lines and 'Role | Engineer' in lines and file.read_bytes()==before
    assert build_profile(lines).skills['imported_skills'][0].name=='Python'

def test_parses_structured_resume_sections():
    profile = build_profile([
        'Professional Summary', 'Quality engineer.', 'Technology Stack',
        'Programming Skill | Python, Java', 'Automation Tools | Selenium, Pytest',
        'KPI’s & Responsibility', 'Improved coverage.', 'Work Experience',
        'Role: SDET', 'Company: Example Corp', 'Duration: Jan 2020 – Present',
        'Location: Bengaluru', 'Responsibilities:', 'Built automation framework.',
        'Continued responsibility across a page boundary.', 'Associations:', 'Example Corp | SDET',
        'Education', 'Bachelor of Engineering | Example University | India',
        'Certifications', 'ISTQB Certified Tester | ISTQB', 'Publications',
        'Inventions and Patents filed for:', 'Test invention | Example Corp',
    ])
    assert [skill.name for skill in profile.skills['programming']] == ['Python', 'Java']
    assert len(profile.experience) == 1
    assert profile.experience[0].company == 'Example Corp'
    assert profile.experience[0].responsibilities[-1].startswith('Continued')
    assert len(profile.achievements) == 1
    assert profile.publications_and_patents == ['Test invention | Example Corp']
