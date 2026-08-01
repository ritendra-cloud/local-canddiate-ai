import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from app.models.candidate import CandidateProfile


HEADING_ALIASES = {
    'professional summary': 'professional_summary',
    'skills': 'skills', 'experience': 'experience', 'achievements': 'achievements',
    'technology stack': 'skills',
    'kpi’s & responsibility': 'achievements',
    "kpi's & responsibility": 'achievements',
    'work experience': 'experience',
    'associations': 'associations',
    'education': 'education',
    'certifications': 'certifications',
    'publications': 'publications_and_patents',
}
SKILL_CATEGORIES = {
    'programming skill': 'programming', 'automation tools': 'automation_testing',
    'data stores': 'databases', 'web servers': 'infrastructure',
    'operating systems': 'operating_systems', 'others': 'tools_and_methodologies',
}


def clean(value: str) -> str:
    return re.sub(r'\s+', ' ', value).strip()


def normalized_heading(value: str) -> str:
    return clean(value).lower().replace('’', "'").rstrip(':').strip()


def extract_docx(path: Path) -> list[str]:
    """Extract ordered paragraph text plus table rows without changing the DOCX."""
    doc = Document(path)
    lines = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    for table in doc.tables:
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells if clean(cell.text)]
            if cells:
                lines.append(' | '.join(cells))
    return lines


def _split_items(value: str) -> list[dict[str, str]]:
    return [{'name': clean(item)} for item in re.split(r'[,;•]+', value) if clean(item)]


def _split_duration(value: str) -> tuple[str | None, str | None]:
    parts = re.split(r'\s+(?:–|—|-)\s+', value, maxsplit=1)
    return (parts[0], parts[1]) if len(parts) == 2 else (value or None, None)


def _candidate_name(lines: list[str]) -> str | None:
    for line in lines:
        if ' | ' not in line:
            continue
        first = clean(line.split('|', 1)[0])
        match = re.match(r'^(.+?)\s+(?:Automation|Manual|Quality)\s+(?:and\s+)?(?:Manual\s+)?Test', first, re.I)
        if match:
            return clean(match.group(1))
    return None


def build_profile(lines: list[str]) -> CandidateProfile:
    sections: dict[str, list[str]] = {'unclassified': []}
    detected: list[str] = []
    current = 'unclassified'
    skills: dict[str, list[dict[str, str]]] = {}

    for line in lines:
        label = normalized_heading(line)
        if label in HEADING_ALIASES:
            current = HEADING_ALIASES[label]
            sections.setdefault(current, [])
            detected.append(current)
            continue
        cells = [clean(cell) for cell in line.split('|')]
        category = SKILL_CATEGORIES.get(normalized_heading(cells[0])) if len(cells) == 2 else None
        if category:
            skills[category] = _split_items(cells[1])
            continue
        # The DOCX contains a header/contact table and an Associations table.
        # Tables are exposed after paragraphs by python-docx, so never attach
        # their three-column rows to the final Publications section.
        if len(cells) == 3 and current not in {'education', 'certifications'}:
            continue
        sections.setdefault(current, []).append(line)

    if not skills and sections.get('skills'):
        skills['imported_skills'] = [item for line in sections['skills'] for item in _split_items(line)]

    experiences: list[dict] = []
    active: dict | None = None
    collecting_responsibilities = False
    for line in sections.get('experience', []):
        field = re.match(r'^(Role|Company|Duration|Location|Responsibilities):\s*(.*)$', line, re.I)
        if field:
            key, value = field.group(1).lower(), clean(field.group(2))
            if key == 'role':
                if active:
                    experiences.append(active)
                active = {'role': value or None, 'responsibilities': [], 'achievements': [], 'technologies': []}
                collecting_responsibilities = False
            elif active is not None:
                if key == 'company': active['company'] = value or None
                elif key == 'location': active['location'] = value or None
                elif key == 'duration': active['start_date'], active['end_date'] = _split_duration(value)
                elif key == 'responsibilities': collecting_responsibilities = True
            continue
        if active is not None and collecting_responsibilities and line:
            active['responsibilities'].append(line)
    if active:
        experiences.append(active)

    education = []
    for line in sections.get('education', []):
        parts = [clean(item) for item in line.split('|')]
        if parts and parts[0]:
            education.append({'degree': parts[0], 'institution': parts[1] if len(parts) > 1 else None})
    certifications = []
    for line in sections.get('certifications', []):
        parts = [clean(item) for item in line.split('|')]
        if parts and parts[0]: certifications.append({'name': parts[0], 'issuer': parts[1] if len(parts) > 1 else None})
    publications = [line for line in sections.get('publications_and_patents', []) if line and not re.match(r'^(inventions?|patents?)\b', line, re.I)]
    summary = '\n'.join(sections.get('professional_summary', [])) or None
    unclassified = [line for line in sections['unclassified'] if line]
    now = datetime.now(timezone.utc).isoformat()
    return CandidateProfile.model_validate({
        'schema_version': '1.0', 'last_updated': now,
        'candidate': {'name': _candidate_name(lines), 'professional_summary': summary},
        'education': education, 'skills': skills, 'experience': experiences,
        'projects': [], 'certifications': certifications,
        'achievements': sections.get('achievements', []), 'publications_and_patents': publications,
        'social_links': {},
        'import_metadata': {'source_type': 'docx', 'imported_at': now,
            'sections_detected': list(dict.fromkeys(detected)),
            'warnings': ['Draft generated deterministically; review all fields before use.'],
            'unclassified_block_count': len(unclassified)},
    })
